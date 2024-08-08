from docx import Document
from docx.shared import Pt  # For font size
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # For handling XML namespaces
import os

def set_borders(cell, border_size=4, border_loc={"left":False, "right":False, "top":False, "bottom":False}):
    tbl = cell._element.getparent().getparent()  # Get the table
    tblPr = tbl.tblPr
    
    # If tblBorders is not present, create it
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    
    # Convert border size from points to twips (1 point = 8 twips)
    border_size_twips = str(border_size * 8)

    # Define and set border properties
    def set_border(name):
        border = tblBorders.find(qn(f'w:{name}'))
        if border is None:
            border = OxmlElement(f'w:{name}')
            tblBorders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size_twips)
    
    if border_loc["left"]: set_border('left')
    if border_loc["right"]: set_border('right')
    if border_loc["top"]: set_border('top')
    if border_loc["bottom"]: set_border('bottom')

def set_font(cell, font_name='Arial', font_size=Pt(12), bold=False):
    # Access the cell text and apply font
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold  # Set bold attribute
        # Ensure that the font name is applied in the document
        rPr = paragraph.runs[0].element.rPr
        rFonts = rPr.rFonts
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)

def center_align_text(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = 1  # 1 represents center alignment

def create_table(doc, name, city):
    
    table = doc.add_table(rows=4, cols=1)
    
    cell11 = table.cell(0, 0)
    cell11.text = "TO:                                       ___ OF ___"
    
    cell21 = table.cell(1, 0)
    cell21.text = name

    cell31 = table.cell(2, 0)
    cell31.text = "IN:"
    
    cell41 = table.cell(3, 0)
    cell41.text = city
    
    # Set font for cells
    set_font(cell11, 'Times New Roman', Pt(24), True)
    set_font(cell21, 'Rockwell', Pt(40), True)
    set_font(cell31, 'Times New Roman', Pt(24), True)
    set_font(cell41, 'Rockwell', Pt(40), True)    

    center_align_text(cell21)
    center_align_text(cell41)
    border_thickness = 4
    set_borders(cell11, border_size=border_thickness, border_loc={"left":True, "right":True, "top":True, "bottom":False}) 
    set_borders(cell21, border_size=border_thickness, border_loc={"left":True, "right":True, "top":False, "bottom":False}) 
    set_borders(cell31, border_size=border_thickness, border_loc={"left":True, "right":True, "top":False, "bottom":False}) 
    set_borders(cell41, border_size=border_thickness, border_loc={"left":True, "right":True, "top":False, "bottom":True}) 

def add_para(doc):
    p_to = doc.add_paragraph()
    run_to = p_to.add_run(' ')
    run_to.font.name = 'Times New Roman'
    run_to.font.size = Pt(12)
def add_name_city(name_city_list, digits, filename):
    result_folder = "results"
    os.makedirs(result_folder, exist_ok=True)
    for k, v in enumerate(name_city_list):
        name = v[0]
        city = v[1]
        create_table(doc, name, city)
        add_para(doc)
        if k%2 != 0:
            doc.add_page_break()
    doc.save(f'{result_folder}/{filename}.docx')   
    # saving digit 
    file_name = f'{result_folder}/{filename}.txt'

    # Write the lines to the file
    with open(file_name, 'w') as file:
        for line in digits:
            file.write(line + '\n')    
doc = Document()
